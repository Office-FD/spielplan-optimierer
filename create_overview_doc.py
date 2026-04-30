"""Erstellt die Uebersichts-DOCX fuer den Spielplan-Optimierer."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path


# ── Farben ────────────────────────────────────────────────────────────────────
BLUE_DARK  = RGBColor(0x1A, 0x3A, 0x5C)   # Titelblau
BLUE_MID   = RGBColor(0x1E, 0x6B, 0xB0)   # Abschnittblau
BLUE_LIGHT = RGBColor(0xD6, 0xE8, 0xF7)   # Tabellenkopf-Hintergrund
GREY_LIGHT = RGBColor(0xF4, 0xF4, 0xF4)   # Zebra-Hintergrund
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
ORANGE     = RGBColor(0xE0, 0x70, 0x20)


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = str(rgb)  # RGBColor.__str__ returns 6-char hex, e.g. '1A3A5C'
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if color:
        run.font.color.rgb = color
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    return p


def add_bullet(doc, text, level=0, bold_prefix=None):
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix + '  ')
        r.bold = True
        r.font.color.rgb = BLUE_MID
    p.add_run(text)
    return p


def add_feature_table(doc, rows, header=('Funktion', 'Beschreibung')):
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Spaltenbreiten
    for cell in table.columns[0].cells:
        cell.width = Cm(5.5)
    for cell in table.columns[1].cells:
        cell.width = Cm(11.0)

    # Kopfzeile
    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = WHITE
        set_cell_bg(hdr[i], BLUE_DARK)

    # Datenzeilen
    for ridx, (label, desc) in enumerate(rows):
        cells = table.rows[ridx + 1].cells
        cells[0].text = label
        cells[1].text = desc
        cells[0].paragraphs[0].runs[0].bold = True
        if ridx % 2 == 0:
            set_cell_bg(cells[0], BLUE_LIGHT)
            set_cell_bg(cells[1], BLUE_LIGHT)
        else:
            set_cell_bg(cells[0], GREY_LIGHT)
            set_cell_bg(cells[1], GREY_LIGHT)

    doc.add_paragraph()


# ── Dokument aufbauen ─────────────────────────────────────────────────────────

doc = Document()

# Seitenränder
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# Standardschrift
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

# ── Titelseite ────────────────────────────────────────────────────────────────

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('Spielplan-Optimierer')
r.font.name  = 'Calibri'
r.font.size  = Pt(28)
r.font.bold  = True
r.font.color.rgb = BLUE_DARK

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('Automatische Spielplanerstellung für Sportligen')
r2.font.size  = Pt(14)
r2.font.color.rgb = BLUE_MID

doc.add_paragraph()

intro = doc.add_paragraph()
intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = intro.add_run('Funktionsübersicht · Stand April 2026')
r3.font.size  = Pt(10)
r3.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
r3.italic = True

doc.add_paragraph()
doc.add_paragraph()

# ── 1. Was ist der Spielplan-Optimierer? ─────────────────────────────────────

add_heading(doc, '1  Was ist der Spielplan-Optimierer?', level=1, color=BLUE_DARK)

p = doc.add_paragraph()
p.add_run(
    'Der Spielplan-Optimierer ist eine Softwarelösung zur automatischen Erstellung '
    'mathematisch optimierter Spielpläne für Sportligen. Er berechnet Spielpläne, '
    'die mehrere konkurrierende Ziele gleichzeitig berücksichtigen: minimale '
    'Reisedistanzen, ausgeglichene Heimspielverteilung und – bei Mehrspartenvereinen – '
    'eine koordinierte Terminplanung über mehrere Ligen hinweg.'
)
p.paragraph_format.space_after = Pt(6)

p2 = doc.add_paragraph()
p2.add_run(
    'Die Lösung richtet sich an Ligaorganisatoren, Verbände und Vereine, die Spielpläne '
    'bisher manuell oder mit einfachen Tabellenkalkulationen erstellen. Der gesamte '
    'Prozess – von der Dateneingabe bis zur fertigen Excel-Datei – wird durch einen '
    'geführten Dialog (Wizard) gesteuert, der keine Programmierkenntnisse voraussetzt.'
)
p2.paragraph_format.space_after = Pt(8)


# ── 2. Kernfunktionen ─────────────────────────────────────────────────────────

add_heading(doc, '2  Kernfunktionen im Überblick', level=1, color=BLUE_DARK)

add_feature_table(doc, [
    ('Mehrere Ligen gleichzeitig',
     'Bis zu 8 Ligen können in einem Durchlauf gemeinsam optimiert werden. '
     'Phase 1 läuft für alle Ligen parallel.'),
    ('Flexible Teamanzahl',
     'Jede Liga kann eine beliebige Teamanzahl haben (mind. 4). '
     'Verschiedene Ligen dürfen unterschiedlich viele Teams haben.'),
    ('Vier Spielmodi',
     'Einfachrunde · Hin-/Rückrunde (Standard) · Dreifachrunde · Turniertag '
     '(2 Spiele pro Team pro Veranstaltung)'),
    ('Distanzoptimierung',
     'Reisedistanzen werden per Google Maps API, CSV/Excel-Datei oder manueller '
     'Eingabe ermittelt und in der Optimierung berücksichtigt.'),
    ('DST-Unterstützung',
     'Doppelspieltage (DST): Zwei aufeinanderfolgende Spieltage werden als Block '
     'behandelt. Im Turniertag-Modus werden alle Spieltage automatisch zu DST-Blöcken.'),
    ('Routing-Optimierung',
     'Bei DST-Blöcken kann optional sichergestellt werden, dass Auswärtsfahrten '
     'nicht über einen konfigurierbaren Umweg-Faktor hinausgehen.'),
    ('Co-Home-Koordination',
     'Mehrspartenvereine (z. B. Damen und Herren im selben Verein) können so geplant '
     'werden, dass ihre Heimspiele möglichst in derselben Kalenderwoche liegen.'),
    ('Pflichtspiele',
     'Bestimmte Begegnungen können auf feste Spieltage und/oder Heimrechte gepinnt '
     'werden (z. B. Eröffnungsspiel, Stadtderby).'),
    ('Heimspiel-Sperrtage',
     'Pro Team können Spieltage gesperrt werden, an denen kein Heimspiel '
     'stattfinden darf (z. B. Hallenbelegung durch andere Veranstaltungen).'),
    ('Konsekutiv-Beschränkung',
     'Maximal 2 aufeinanderfolgende Heim- oder Auswärtsspiele; bei DST-Beteiligung '
     'bis zu 3 erlaubt. Nie mehr als 3 in Folge.'),
    ('Gewichtete Optimierung',
     'Alle Optimierungsziele sind individuell pro Liga auf einer Skala 0–10 '
     'gewichtbar: Reisedistanz, Reisegerechtigkeit, Heimrechtswechsel, Co-Home-Bonus.'),
    ('Excel-Ausgabe',
     'Jede Liga wird als farblich aufbereitete Excel-Datei exportiert. '
     'Bei Mehrspartenvereinen wird zusätzlich eine Co-Home-Übersichtsdatei erstellt.'),
    ('Kalenderintegration',
     'Ein Rahmenterminplan (Excel) kann geladen werden, um Spieltage automatisch '
     'Kalenderwochen und Datumsbereichen zuzuordnen.'),
    ('Mehrfachdurchläufe',
     'Jede Liga wird mit mehreren Zufalls-Seeds gelöst; die beste Lösung gewinnt. '
     'Konfigurierbar: 1–4 Seeds.'),
    ('SA-Nachbearbeitung',
     'Nach der CP-SAT-Optimierung verbessert ein Simulated-Annealing-Algorithmus '
     'die Heimrechtzuweisungen weiter (typisch: 3–8 % weniger Gesamtkilometer).'),
])


# ── 3. Spielmodi ──────────────────────────────────────────────────────────────

add_heading(doc, '3  Spielmodi', level=1, color=BLUE_DARK)

add_feature_table(doc, [
    ('Einfachrunde',
     'Jede Paarung spielt genau einmal. Die Heimrechtzuweisung wird vom Optimierer '
     'frei gewählt, um die Gesamtkilometer zu minimieren.'),
    ('Hin-/Rückrunde',
     'Standard-Modus. Jede Paarung spielt zweimal: einmal mit Heimrecht A und '
     'einmal mit Heimrecht B. Die Zuordnung, wer in welcher Runde zu Hause spielt, '
     'wird optimiert.'),
    ('Dreifachrunde',
     'Jede Paarung spielt dreimal. Ein Team erhält 2 Heimspiele, das andere 1 – '
     'wer was bekommt, entscheidet der Optimierer nach Distanzkriterien.'),
    ('Turniertag',
     'Hin-/Rückrunde, bei der alle Spieltage paarweise als DST-Blöcke behandelt '
     'werden. Jedes Team spielt pro Veranstaltungstag 2 Spiele. Geeignet für '
     'Sportarten, die mehrere Partien an einem Tag austragen.'),
], header=('Modus', 'Beschreibung'))


# ── 4. Optimierungsprozess ────────────────────────────────────────────────────

add_heading(doc, '4  Dreiphasiger Optimierungsprozess', level=1, color=BLUE_DARK)

add_heading(doc, 'Phase 1 – Unabhängige Liga-Optimierung (parallel)', level=2, color=BLUE_MID)
p = doc.add_paragraph(
    'Alle Ligen werden gleichzeitig und vollständig unabhängig voneinander optimiert. '
    'Der CP-SAT-Solver (Google OR-Tools) berechnet für jede Liga den besten Spielplan '
    'unter den individuellen Constraints. Da die Ligen in Phase 1 nichts voneinander '
    'wissen, läuft dieser Schritt vollständig parallel und skaliert mit der CPU-Anzahl.'
)
p.paragraph_format.space_after = Pt(6)

add_heading(doc, 'Phase 2 – Kombiniertes Modell (Co-Home-Koordination)', level=2, color=BLUE_MID)
p = doc.add_paragraph(
    'Alle Ligen werden in einem gemeinsamen Modell zusammengeführt. Der Co-Home-Bonus '
    'sorgt dafür, dass Mehrspartenvereine ihre Heimspiele möglichst in denselben '
    'Kalenderwochen haben. Die Lösungen aus Phase 1 dienen als Warm-Start-Hints, '
    'was die Konvergenz erheblich beschleunigt. Eine Ligahierarchie steuert, welche '
    'Ligen bei Konflikten Priorität erhalten.'
)
p.paragraph_format.space_after = Pt(6)

add_heading(doc, 'Phase 3 – SA-Nachbearbeitung (Heimrecht-Feinoptimierung)', level=2, color=BLUE_MID)
p = doc.add_paragraph(
    'Ein Simulated-Annealing-Algorithmus optimiert die Heimrechtzuweisungen weiter, '
    'ohne das Termingerüst zu verändern. Alle Hard-Constraints (Pflichtspiele, '
    'Sperrtage, DST-Bindungen) bleiben garantiert erhalten. Typische Verbesserung: '
    '3–8 % weniger Gesamtkilometer.'
)
p.paragraph_format.space_after = Pt(8)


# ── 5. Eingabe-Konfiguration ──────────────────────────────────────────────────

add_heading(doc, '5  Eingabe-Konfiguration (geführter Wizard)', level=1, color=BLUE_DARK)

p = doc.add_paragraph(
    'Die gesamte Konfiguration erfolgt über eine browserbasierte Streamlit-Oberfläche '
    'in 9 Schritten (Schritt 0–8). Die Anwendung wird lokal gestartet und im Browser '
    'bedient; technische Kenntnisse sind nicht erforderlich.'
)
p.paragraph_format.space_after = Pt(6)

steps = [
    ('Schritt 0', 'Ligen & Teams',
     'Anzahl der Ligen, Liga-ID, Name, Teams, Standorte, Spielmodus, Hierarchiegewicht. '
     'Excel-Vorlage zum Laden und Speichern der Konfiguration.'),
    ('Schritt 1', 'Distanzmatrizen',
     'Auswahl der Methode: Google Maps API (automatisch) · CSV/Excel-Datei · manuelle Eingabe. '
     'Fehleranzeige bei ungültigen Adressen; Adressvorschau vor dem API-Aufruf.'),
    ('Schritt 2', 'Kalender & DST',
     'Optionaler Rahmenterminplan aus Excel; manuelle DST-Block-Konfiguration als Fallback.'),
    ('Schritt 3', 'Routing & Gewichte',
     'DST-Routing-Beschränkung (max. Umweg in Prozent) und Optimierungsziele gewichten (0–10): '
     'Reisedistanz, Reisegerechtigkeit, Heimrechtswechsel, Co-Home-Bonus.'),
    ('Schritt 4', 'Pflichtspiele',
     'Bestimmte Begegnungen auf feste Spieltage und/oder Heimrechte festlegen.'),
    ('Schritt 5', 'Sperrtage',
     'Pro Team: Spieltage sperren, an denen kein Heimspiel erlaubt ist.'),
    ('Schritt 6', 'Co-Home',
     'Mehrspartenvereine definieren: welche Teams in welchen Ligen zum selben Verein gehören.'),
    ('Schritt 7', 'Solver',
     'Zeitlimits für Phase 1/2/3, Anzahl Seeds, Intensitätsmodus (Standard/Intensiv/Nachtlauf).'),
    ('Schritt 8', 'Optimierung & Ergebnisse',
     'Konfigurationsübersicht, Optimierungslauf starten, Ergebnisse anzeigen, '
     'Excel-Dateien herunterladen.'),
]

table = doc.add_table(rows=1 + len(steps), cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.LEFT
for w, col in zip([Cm(2.2), Cm(4.2), Cm(10.1)], table.columns):
    for cell in col.cells:
        cell.width = w

hdr = table.rows[0].cells
for i, h in enumerate(['Schritt', 'Thema', 'Inhalt']):
    hdr[i].text = h
    hdr[i].paragraphs[0].runs[0].bold = True
    hdr[i].paragraphs[0].runs[0].font.color.rgb = WHITE
    set_cell_bg(hdr[i], BLUE_DARK)

for ridx, (num, title_s, desc) in enumerate(steps):
    cells = table.rows[ridx + 1].cells
    cells[0].text = num
    cells[1].text = title_s
    cells[2].text = desc
    cells[1].paragraphs[0].runs[0].bold = True
    bg = BLUE_LIGHT if ridx % 2 == 0 else GREY_LIGHT
    for c in cells:
        set_cell_bg(c, bg)

doc.add_paragraph()


# ── 6. Technische Basis ───────────────────────────────────────────────────────

add_heading(doc, '6  Technische Basis', level=1, color=BLUE_DARK)

add_feature_table(doc, [
    ('Optimierungsverfahren',
     'CP-SAT (Constraint Programming / SAT) via Google OR-Tools – '
     'garantiert optimale oder nachweislich nahezu-optimale Lösungen.'),
    ('Nachbearbeitung',
     'Simulated Annealing (SA) für Heimrecht-Feinoptimierung nach dem CP-SAT-Lauf.'),
    ('Parallelisierung',
     'Phase 1: mehrere Ligen und mehrere Seeds laufen gleichzeitig '
     '(ProcessPoolExecutor, bis zu 8 CPU-Kerne nutzbar).'),
    ('Benutzeroberfläche',
     'Streamlit – läuft lokal im Browser, Start per start.bat oder '
     '„streamlit run app.py". Keine Installation eines separaten Clients nötig.'),
    ('Programmiersprache',
     'Python 3.10+. Benötigte Pakete: streamlit, ortools, openpyxl, numpy, pandas, python-docx.'),
    ('Ausgabeformat',
     'Excel (.xlsx) mit farblich codierten Spielplänen und Statistikblättern; '
     'optionale Co-Home-Übersichtsdatei.'),
    ('Skalierbarkeit',
     'Bis zu 8 Ligen gleichzeitig; bis zu 20+ Teams pro Liga; '
     'Warnung bei mehr als 48 Teams gesamt (Phase 2 dauert dann länger).'),
    ('Laufzeitrichtwerte',
     'Phase 1: 15–60 min (parallel). Phase 2: 30–120 min je nach Teamanzahl. '
     'Phase 3: konfigurierbar, Standard 2 min pro Liga.'),
], header=('Merkmal', 'Details'))


# ── 7. Anwendungsszenarien ────────────────────────────────────────────────────

add_heading(doc, '7  Beispiel-Anwendungsszenarien', level=1, color=BLUE_DARK)

scenarios = [
    ('Regionaler Verband – 4 Ligen',
     '4 Ligen (z. B. 1. BL, 2. BL, Damen, Junioren) mit je 10–12 Teams. '
     'Co-Home-Koordination für Mehrspartenvereine. Rahmenterminplan aus Excel. '
     'Gesamtlaufzeit ca. 2–3 Stunden (Nachtlauf).'),
    ('Kleiner Verband – 1–2 Ligen',
     '1 oder 2 Ligen mit 6–8 Teams. Distanzen manuell eingegeben. '
     'Keine Co-Home-Koordination. Laufzeit unter 15 Minuten.'),
    ('Turnierserie',
     'Turniertag-Modus: mehrere Teams spielen an einem Wochenende je 2 Partien. '
     'Alle Spieltage automatisch als DST-Blöcke. Routing-Optimierung aktiv.'),
    ('Dreifachrunde',
     '3 Begegnungen pro Paarung. Eine Seite erhält automatisch 2 Heimspiele, '
     'die andere 1 – Zuteilung erfolgt reisedistanzoptimiert.'),
]

for name, desc in scenarios:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(name + ':  ')
    r.bold = True
    r.font.color.rgb = BLUE_MID
    p.add_run(desc)


# ── 8. Abgrenzung / Was das Tool nicht macht ──────────────────────────────────

add_heading(doc, '8  Abgrenzung', level=1, color=BLUE_DARK)

p = doc.add_paragraph(
    'Der Spielplan-Optimierer erstellt Spielpläne auf Spieltag-Ebene: er legt fest, '
    'welche Begegnung an welchem Spieltag stattfindet und welches Team Heimrecht hat. '
    'Folgendes ist nicht Bestandteil des aktuellen Funktionsumfangs:'
)
p.paragraph_format.space_after = Pt(4)

not_included = [
    'Konkrete Uhrzeiten oder Hallenbelegungspläne',
    'Schiedsrichter-Ansetzungen',
    'Ligaverwaltung oder Ergebniserfassung',
    'Cloud-Betrieb oder zentraler Web-Zugang – die Anwendung läuft lokal auf dem eigenen Rechner',
]
for item in not_included:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.add_run(item)

doc.add_paragraph()


# ── Fußzeile ──────────────────────────────────────────────────────────────────

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Spielplan-Optimierer · Floorball Verband Deutschland e. V. · April 2026')
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
r.italic = True


# ── Speichern ─────────────────────────────────────────────────────────────────

out = Path(__file__).parent / 'Spielplan-Optimierer_Uebersicht.docx'
doc.save(out)
print(f'Gespeichert: {out}')
